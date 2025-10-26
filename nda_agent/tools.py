"""Utility tools for the NDA reviewer workflow."""
from __future__ import annotations

import json
import logging
import os
import tempfile
import time
from pathlib import Path
from typing import Dict, Iterable, Optional

import requests
import gspread
from google.oauth2.service_account import Credentials
from pypdf import PdfReader, PdfWriter
from pypdf.annotations import AnnotationBuilder
from docx import Document  # type: ignore

LOGGER = logging.getLogger(__name__)

REQUIRED_COLUMNS: Iterable[str] = (
    "legal_name",
    "address",
    "signer_name",
    "title",
    "email",
    "phone",
    "effective_date",
)

PDF_FIELD_MAP = {
    "CompanyName": "legal_name",
    "Address": "address",
    "SignerName": "signer_name",
    "Title": "title",
    "Email": "email",
    "Phone": "phone",
    "EffectiveDate": "effective_date",
}


class ToolError(RuntimeError):
    """Base error class for tool failures."""


def _load_service_account_credentials(path: str) -> Credentials:
    LOGGER.debug("Loading Google service account credentials from %s", path)
    with open(path, "r", encoding="utf-8") as handle:
        data = json.load(handle)
    scopes = [
        "https://www.googleapis.com/auth/spreadsheets.readonly",
        "https://www.googleapis.com/auth/drive.readonly",
    ]
    return Credentials.from_service_account_info(data, scopes=scopes)


def fetch_profile_from_sheets(sheet_id: str, *, worksheet: Optional[str] = None) -> Dict[str, str]:
    """Fetch the first row from a Google Sheet as a profile dict."""
    sa_path = os.getenv("GOOGLE_SHEETS_SA_JSON")
    if not sa_path:
        raise ToolError("GOOGLE_SHEETS_SA_JSON is not configured.")
    creds = _load_service_account_credentials(sa_path)
    client = gspread.authorize(creds)
    spreadsheet = client.open_by_key(sheet_id)
    if worksheet:
        sheet = spreadsheet.worksheet(worksheet)
    else:
        sheet = spreadsheet.sheet1

    rows = sheet.get_all_records(head=1)
    if not rows:
        raise ToolError("No rows found in the Google Sheet.")
    profile = rows[0]

    missing = [col for col in REQUIRED_COLUMNS if col not in profile]
    if missing:
        raise ToolError(f"Missing required columns in sheet: {', '.join(missing)}")

    return {key: str(profile.get(key, "")).strip() for key in profile}


def _create_cover_page(profile: Dict[str, str]) -> Path:
    writer = PdfWriter()
    page = writer.add_blank_page(width=612, height=792)
    y = 760
    lines = ["Buyer Profile"] + [f"{key}: {value}" for key, value in profile.items()]
    for line in lines:
        annotation = AnnotationBuilder.free_text(
            text=line,
            rect=(40, y - 12, 580, y + 12),
            font="Helvetica",
            font_size=12,
        )
        page.add_annotation(annotation)
        y -= 24

    temp = Path(tempfile.mkstemp(suffix="_cover.pdf")[1])
    with open(temp, "wb") as handle:
        writer.write(handle)
    return temp


def _fill_pdf_form(input_path: Path, output_path: Path, profile: Dict[str, str]) -> None:
    reader = PdfReader(str(input_path))
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)

    form_fields = reader.get_form_text_fields()
    if form_fields:
        LOGGER.info("Found form fields: %s", list(form_fields.keys()))
        for field_name, sheet_key in PDF_FIELD_MAP.items():
            if field_name in form_fields:
                value = profile.get(sheet_key, "")
                for page in writer.pages:
                    writer.update_page_form_field_values(page, {field_name: value})
    else:
        LOGGER.info("No form fields detected; prepending generated cover page")
        cover_path = _create_cover_page({k: profile.get(k, "") for k in REQUIRED_COLUMNS})
        cover_reader = PdfReader(str(cover_path))
        combined_writer = PdfWriter()
        for page in cover_reader.pages:
            combined_writer.add_page(page)
        for page in writer.pages:
            combined_writer.add_page(page)
        writer = combined_writer

    with open(output_path, "wb") as handle:
        writer.write(handle)


def _fill_docx(input_path: Path, output_path: Path, profile: Dict[str, str]) -> None:
    doc = Document(str(input_path))
    replacements = {f"{{{{{key}}}}}": value for key, value in profile.items()}

    def replace_text(text: str) -> str:
        for placeholder, value in replacements.items():
            if placeholder in text:
                text = text.replace(placeholder, value)
        return text

    for paragraph in doc.paragraphs:
        paragraph.text = replace_text(paragraph.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = replace_text(cell.text)

    doc.save(output_path)


def fill_pdf_or_docx(input_file: str, profile: Dict[str, str], *, output_dir: Optional[str] = None) -> str:
    """Fill a PDF/DOCX document with profile details.

    Returns the path to the generated document (PDF for PDFs, DOCX for word files).
    """
    input_path = Path(input_file)
    if not input_path.exists():
        raise ToolError(f"Input file not found: {input_file}")

    output_dir_path = Path(output_dir or input_path.parent)
    output_dir_path.mkdir(parents=True, exist_ok=True)

    if input_path.suffix.lower() == ".pdf":
        output_path = output_dir_path / f"{input_path.stem}.filled.pdf"
        _fill_pdf_form(input_path, output_path, profile)
        return str(output_path)

    if input_path.suffix.lower() == ".docx":
        output_path = output_dir_path / f"{input_path.stem}.filled.docx"
        _fill_docx(input_path, output_path, profile)
        return str(output_path)

    raise ToolError("Unsupported file type. Only PDF and DOCX files are supported.")


def send_webhook(url: str, payload: Dict[str, object], *, retries: int = 2, timeout: int = 5) -> bool:
    """POST the payload to the webhook URL with simple retry logic."""
    if not url:
        LOGGER.info("Webhook URL not provided; skipping call")
        return False

    for attempt in range(1, retries + 2):
        try:
            LOGGER.info("Sending webhook attempt %s to %s", attempt, url)
            response = requests.post(url, json=payload, timeout=timeout)
            if response.ok:
                LOGGER.info("Webhook delivered successfully with status %s", response.status_code)
                return True
            LOGGER.warning("Webhook responded with %s: %s", response.status_code, response.text)
        except requests.RequestException as exc:
            LOGGER.error("Webhook error: %s", exc)
        time.sleep(1)
    return False


__all__ = [
    "fetch_profile_from_sheets",
    "fill_pdf_or_docx",
    "send_webhook",
    "ToolError",
    "REQUIRED_COLUMNS",
    "PDF_FIELD_MAP",
]
